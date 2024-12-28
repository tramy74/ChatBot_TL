import os
import pdfplumber
import docx
from sqlalchemy.orm import Session
from app.models import Document, Embedding
import openai
from app.utils.image_processing import segment_and_extract_text_from_image
from app.utils.video_transcription import transcribe_video_audio

UPLOAD_FOLDER = "uploaded_files"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
openai.api_key = os.getenv("OPENAI_API_KEY")

def upload_and_process_file(file, title, author, doc_type, summary, db: Session):
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    with open(file_path, "wb") as f:
        f.write(file.file.read())

    if doc_type == "PDF":
        content = extract_text_from_pdf(file_path)
    elif doc_type == "DOCX":
        content = extract_text_from_docx(file_path)
    elif doc_type == "Image":
        content = segment_and_extract_text_from_image(file_path)
    elif doc_type == "Video":
        content = transcribe_video_audio(file_path)
    else:
        content = None

    db_document = Document(
        title=title,
        author=author,
        doc_type=doc_type,
        summary=summary,
        content_url=file_path
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)

    if content:
        embedding = generate_embedding(content)
        db_embedding = Embedding(
            document_id=db_document.document_id,
            embedding=embedding
        )
        db.add(db_embedding)
        db.commit()

    return {"message": "File uploaded and processed successfully", "document_id": db_document.document_id}

def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        return "".join(page.extract_text() for page in pdf.pages)

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)

def generate_embedding(text):
    response = openai.Embedding.create(input=text, model="text-embedding-ada-002")
    return response['data'][0]['embedding']
